"""Create DOCX document files for download.

Tier: std-create (Tier 1) — single-shot file generation.

**std-create**: one function call, one file out. Used when the agent has
data in memory and just needs a polished .docx for download (KPI
reports, exports, ad-hoc summaries). Auto-loaded for every agent that
pulls in `agentino.tools.std`.

For multi-step document work — open existing files, fill templates with
variables, edit specific paragraphs, branded deliverables — the
convention is **Tier 2 / `skill-docx`**: a per-tenant skill bundle
(`<tenant>/skills/docx/` containing `SKILL.md` + tools + templates),
opt-in via `skills:` in that tenant's agents.yml. Tier 2 isn't
centralized — each tenant ships only what its agents need. If/when 2+
tenants need the same skill, lift it to a peer `agentic/skills/` dir.

Tool name in agents.yml: `create_document` (unchanged for back-compat).
"""

import io
import re
from datetime import datetime

from agentino.core.tool import tool

from .storage import get_default_store


def _setup_styles(doc):
    """Configure document styles for professional appearance."""
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Pt, RGBColor

    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"
    style.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.15

    for level, size in [(0, 26), (1, 22), (2, 16), (3, 13)]:
        name = f"Heading {level}" if level > 0 else "Title"
        if name in doc.styles:
            hs = doc.styles[name]
            hs.font.name = "Calibri"
            hs.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
            hs.font.size = Pt(size)
            if level > 0:
                hs.paragraph_format.space_before = Pt(16)
                hs.paragraph_format.space_after = Pt(6)

    for list_style in ["List Bullet", "List Number"]:
        if list_style in doc.styles:
            ls = doc.styles[list_style]
            ls.font.size = Pt(11)
            ls.font.name = "Calibri"
            ls.paragraph_format.space_after = Pt(3)


def _parse_markdown_to_docx(doc, content: str):
    """Parse simple markdown into python-docx elements."""
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Headers
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        # Bullet list
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_formatting(p, text)
            i += 1
            continue

        # Numbered list
        m = re.match(r"^\d+\.\s+(.+)", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_inline_formatting(p, m.group(1))
            i += 1
            continue

        # Table — detect markdown table (lines with |)
        if (
            "|" in stripped
            and i + 1 < len(lines)
            and re.match(r"^\s*\|[-:|]+\|", lines[i + 1].strip())
        ):
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            _add_table(doc, table_lines)
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_inline_formatting(p, stripped)
        i += 1


def _add_inline_formatting(paragraph, text: str):
    """Handle **bold** and *italic* inline formatting."""
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _add_table(doc, lines: list[str]):
    """Parse markdown table lines into a styled docx table."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    from docx.shared import Pt, RGBColor

    headers = [c.strip() for c in lines[0].strip("|").split("|")]
    data_rows = []
    for line in lines[2:]:
        cells = [c.strip() for c in line.strip("|").split("|")]
        data_rows.append(cells)

    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.style = "Table Grid"

    # Style header row
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        # Gray background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F3F4F6"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Style data rows with zebra striping
    for ri, row in enumerate(data_rows):
        for ci, val in enumerate(row):
            if ci < len(table.columns):
                cell = table.rows[ri + 1].cells[ci]
                p = cell.paragraphs[0]
                p.clear()
                run = p.add_run(val)
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
                # Zebra stripe
                if ri % 2 == 1:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F9FAFB"/>')
                    cell._tc.get_or_add_tcPr().append(shading)


@tool
async def create_document(title: str, content: str, filename: str = "") -> str:
    """Create a Word (.docx) document and attach it for download. Supports markdown formatting: headers (#, ##, ###), bold (**text**), italic (*text*), bullet lists (- item), numbered lists (1. item), and markdown tables.

    Args:
        title: Document title (appears as heading in the document)
        content: Document content in markdown format
        filename: Optional output filename (without extension). Defaults to title-based name.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    _setup_styles(doc)

    # Set narrow margins for more content area
    for section in doc.sections:
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(54)

    doc.add_heading(title, level=0)

    now = datetime.now().strftime("%B %d, %Y")
    subtitle = doc.add_paragraph(f"Generated {now}")
    subtitle.style = doc.styles["Subtitle"]

    _parse_markdown_to_docx(doc, content)

    if not filename:
        safe = title.lower().replace(" ", "-").replace("/", "-")[:40]
        filename = safe
    now_short = datetime.now().strftime("%Y-%m-%d")
    final_filename = f"{filename}-{now_short}.docx"

    buf = io.BytesIO()
    doc.save(buf)
    content_bytes = buf.getvalue()

    # Source markdown is the cleanest indexable text — strip code fences.
    extracted = re.sub(r"```[a-z]*\n.*?```", "", content, flags=re.DOTALL)[:50_000]

    try:
        stored = get_default_store().save(
            content_bytes=content_bytes,
            filename=final_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            title=title,
            extracted_text=extracted,
        )
    except Exception as e:
        return f"Document generation failed: {e}"

    return f"Document created. Download: [{stored.filename}]({stored.url})"
